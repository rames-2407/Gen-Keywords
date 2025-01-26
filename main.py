import os
from flask import Flask, jsonify, render_template, request
from werkzeug.utils import secure_filename
import pdfplumber  # type: ignore
from docx import Document  # type: ignore
import re
import PyPDF2  # type: ignore
from openai import OpenAI
import chardet  # type: ignore
import win32com.client as win32
import pythoncom
from fuzzywuzzy import fuzz # type: ignore
from dotenv import load_dotenv

load_dotenv()


# OpenAI client configuration
client = OpenAI(
    base_url=os.getenv("OPENAI_API_URL"),
    api_key= os.getenv("OPENAI_API_KEY"),
)

app = Flask(__name__)
UPLOAD_FOLDER = './uploads'
ALLOWED_EXTENSIONS = {'docx', 'pdf', 'txt','doc'}
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


def extract_sentences_from_pdf(pdf_path):
    """Extract all sentences from a PDF using regex."""
    if not os.path.exists(pdf_path):
        print("File not found. Check the file path.")
        return []

    all_sentences = []
    with pdfplumber.open(pdf_path) as pdf:
   
        full_text = "\n".join(page.extract_text() for page in pdf.pages if page.extract_text())

   
    sentence_pattern = r'\d+\.\s(?:[A-Za-z0-9\s,\.&\-]+(?:\([0-9]{4}[a-z]?\))?\.\s)([^\.]+(?:\s[^\.]+)*)'
    sentences = re.findall(sentence_pattern, full_text)

    all_sentences = [sentence.replace('\n', ' ').strip() for sentence in sentences if sentence.strip()]
    return all_sentences

def generate_keywords_from_sentence(sentences):
    """Generate keywords for a given sentence using the OpenAI API."""
    completion = client.chat.completions.create(
        model="mistralai/mixtral-8x7b-instruct-v0.1",
        messages=[{"role": "user", "content": f"Generate keywords for the following sentence: '{sentences}'"}],
        temperature=0.2,
        top_p=0.7,  
        max_tokens=1024,
        stream=True
    )

    keywords = ""
    for chunk in completion:
        if chunk.choices[0].delta.content is not None:
            keywords += chunk.choices[0].delta.content

    return keywords.strip()


def classify_keywords(keywords):
    """Classify keywords into Primary and Secondary."""

    keyword_list = [kw.strip() for kw in keywords.split(",") if kw.strip()]

    num_primary = max(1, len(keyword_list) // 2)  
    primary_keywords = keyword_list[:num_primary]
    secondary_keywords = keyword_list[num_primary:]

    return primary_keywords, secondary_keywords


def allowed_file(filename):
    """Check if the file has an allowed extension."""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def convert_doc_to_docx(doc_path):
    """Convert .doc file to .docx format."""
    if not os.path.exists(doc_path):
        raise FileNotFoundError(f"The file {doc_path} does not exist.")
    docx_path = doc_path.replace(".doc", ".docx")
    pythoncom.CoInitialize()
    word = win32.Dispatch("Word.Application")
    word.Visible = False
    try:
        doc = word.Documents.Open(os.path.abspath(doc_path))
        doc.SaveAs(os.path.abspath(docx_path), FileFormat=16) 
        doc.Close()
    finally:
        word.Quit()
    return docx_path

def extract_keywords_section(docx_path):
    """Extract the 'Keywords' section from a DOCX file."""
    doc = Document(docx_path)
    keywords = []

    # Check paragraphs for keywords
    for para in doc.paragraphs:
        text = para.text.strip()
        if re.search(r'(?i)\bkeywords?\b', text):  # Match "Keywords" section
            match = re.search(r'(?i)\bkeywords?\b\s*[:\-\–]?\s*(.*)', text)
            if match:
                keywords.extend([kw.strip() for kw in re.split(r'[,:;]', match.group(1)) if kw.strip()])

    # Check tables for keywords if not found in paragraphs
    if not keywords:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if re.search(r'(?i)\bkeywords?\b', text):
                        match = re.search(r'(?i)\bkeywords?\b\s*[:\-\–]?\s*(.*)', text)
                        if match:
                            keywords.extend([kw.strip() for kw in re.split(r'[,:;]', match.group(1)) if kw.strip()])
    
    return ", ".join(keywords) if keywords else "No keywords found"


def extract_text(file_path, file_extension):
    """Extract text from files based on their extension."""
    if file_extension == 'docx':
        return extract_text_from_docx(file_path)
    elif file_extension == 'pdf':
        return extract_text_from_pdf(file_path)
    elif file_extension == 'txt':
        return extract_text_from_txt(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_extension}")

def extract_text_from_docx(docx_path):
    """Extract text from a DOCX file."""
    doc = Document(docx_path)
    return [para.text for para in doc.paragraphs]

def extract_text_from_pdf(pdf_path):
    """Extract text from a PDF file."""
    text = []
    with open(pdf_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        for page in reader.pages:
            text.append(page.extract_text())
    return text

def extract_text_from_txt(txt_path):
    """Extract text from a TXT file with encoding detection."""
    with open(txt_path, 'rb') as file:
        raw_data = file.read()
        detected = chardet.detect(raw_data)
        text = raw_data.decode(detected['encoding'])
    return text.splitlines()

def extract_keywords_section_from_text(full_text):
    """Extract keywords from the text by identifying 'Keywords' sections."""
    keywords = []
    for para in full_text:
        match = re.search(r'(?i)\bkeyword[s]?\b[^\n]*[:\-\–]?\s*(.*?)(?:\n|$)', para)
        if match:
            keywords_line = match.group(1)
            keywords.extend([kw.strip() for kw in re.split(r'[,:;]', keywords_line) if kw.strip()])

    keywords=list(set(keywords))
    return " , ".join(keywords) 

def find_fuzzy_matches(file_1_keywords, file_2_keywords, threshold=90):
    """
    Find fuzzy matches between keywords from two files using a similarity threshold.
    """
    matched_keywords = []

    for kw1 in file_1_keywords:
        for kw2 in file_2_keywords:
            if fuzz.ratio(kw1['keyword'].strip().lower(), kw2['keyword'].strip().lower()) >= threshold:
                matched_keywords.append({
                    'keyword': kw1['keyword'],
                    'source_file': kw2['source_file'],  
                    'type': kw1['type']  # Primary or Secondary Keyword for kw1
                })

    return matched_keywords

@app.route('/', methods=['GET', 'POST'])
def home():
    result = None
    error = None
    try:
        if request.method == 'POST':
          
            primary_files = request.files.getlist('file1')
            if not primary_files or all(file.filename == '' for file in primary_files):
                raise ValueError("At least one primary file is required.")
            if not secondary_files:
                raise ValueError("At least one secondary file is required.")
          
            result = {'file_1': []}
            file_1_keywords = []

            for primary_file in primary_files:
                if primary_file.filename == '':
                    continue
                if not allowed_file(primary_file.filename):
                    raise ValueError(f"Unsupported file type for {primary_file.filename}. Please upload DOCX, PDF, or TXT files.")
                
              
                filename_1 = secure_filename(primary_file.filename)
                file_path_1 = os.path.join(app.config['UPLOAD_FOLDER'], filename_1)
                primary_file.save(file_path_1)

              
                bold_sentences_1 = extract_sentences_from_pdf(file_path_1)

             
                all_data_1 = []
                for sentence in bold_sentences_1:
                    keywords = generate_keywords_from_sentence(sentence)
                    keyword_list = keywords.split("\n")
                    primary_keywords, secondary_keywords = classify_keywords(", ".join(keyword_list))
                    all_data_1.append({
                        'sentence': sentence,
                        'primary_keywords': primary_keywords,
                        'secondary_keywords': secondary_keywords
                    })

                for data in all_data_1:
                    file_1_keywords.extend([
                        {'keyword': kw, 'source_file': filename_1, 'type': 'Primary Keyword'}
                        for kw in data['primary_keywords']
                    ])
                    file_1_keywords.extend([
                        {'keyword': kw, 'source_file': filename_1, 'type': 'Secondary Keyword'}
                        for kw in data['secondary_keywords']
                    ])

                result['file_1'].append({'filename': filename_1, 'data': all_data_1})

            secondary_files = request.files.getlist('file2')
            if secondary_files:
                matching_keywords = []
                for secondary_file in secondary_files:
                    if secondary_file.filename == '':
                        continue
                    if not allowed_file(secondary_file.filename):
                        raise ValueError(f"Unsupported file type for {secondary_file.filename}. Please upload DOCX, PDF, or TXT files.")
                    
   
                    filename_2 = secure_filename(secondary_file.filename)
                    file_path_2 = os.path.join(app.config['UPLOAD_FOLDER'], filename_2)
                    secondary_file.save(file_path_2)

                    file_extension_2 = filename_2.rsplit('.', 1)[1].lower()
                    if file_extension_2 == "doc":
                        file_path_2 = convert_doc_to_docx(file_path_2)
                        file_extension_2 = "docx"

                    extracted_keywords_2 = extract_keywords_section(file_path_2)
                    file_2_keywords = [{'keyword': kw.strip(), 'source_file': filename_2} for kw in extracted_keywords_2.split(",") if kw.strip()]

                    matches = find_fuzzy_matches(file_1_keywords, file_2_keywords)
                    matching_keywords.extend(matches)

                   
                    if 'file_2' not in result:
                        result['file_2'] = []
                    result['file_2'].append({
                        'filename': filename_2,
                        'keywords': [kw['keyword'] for kw in file_2_keywords]
                    })

        
                result['matching_keywords'] = matching_keywords

    except Exception as e:
        error = f"Unexpected error: {e}"

    return render_template('index.html', result=result, error=error)


if __name__ == "__main__":
    app.run(debug=True, use_reloader=False)

